# -*- coding: utf-8 -*-
"""
Solray Mobile – Solar proposal app.
Generates one Word file: page 1 = opening letter (from "לכבוד דייר.docx"),
page 2 = price quote (from "פורמט הצעת מחיר.docx"), merged with a page break.

Form inputs: representative, client_name, address, city, roof_sqm, price_per_kw.
Output saved to: <this folder>/לקוחות/<client_name>/
"""
import json
import math
import os
import re
from datetime import datetime

SOLRAY_ROOT = os.path.dirname(os.path.abspath(__file__))
CONFIG_PATH = os.path.join(SOLRAY_ROOT, "config.json")
OUTPUT_DIR = os.environ.get('PERSISTENT_STORAGE_PATH', os.path.join(os.getcwd(), 'לקוחות'))

# יצירת התיקייה רק אם זה לא הנתיב של הדיסק הקבוע (שנוצר אוטומטית)
if not os.path.exists(OUTPUT_DIR):
    try:
        os.makedirs(OUTPUT_DIR, exist_ok=True)
    except OSError as e:
        print(f"Note: Could not create directory {OUTPUT_DIR}: {e}")
TEMPLATE_LETTER = "לכבוד דייר.docx"
TEMPLATE_QUOTE = "פורמט הצעת מחיר - מעוצב.docx"
LOGO_NAMES = ("לוגו.png", "לוגו.jpg")

REPRESENTATIVES = [
    {"full": "עידן עמר", "first": "עידן"},
    {"full": "נתנאל עמר", "first": "נתנאל"},
    {"full": "רון הלל", "first": "רון"},
]

CITIES = [
    "אום אל-פחם", "אופקים", "אור יהודה", "אילת", "אלעד", "אשדוד", "אשקלון", "באקה אל גרביה",
    "באר יעקב", "באר שבע", "בית שמש", "בני ברק", "בת ים", "גבעתיים", "דאלית אל-כרמל", "דימונה",
    "הוד השרון", "הרצליה", "חדרה", "חולון", "חיפה", "טבריה", "טייבה", "טירת כרמל", "טמרה",
    "יבנה", "יפו", "יהוד-מונוסון", "ירושלים", "כפר יונה", "כפר סבא", "כרמיאל", "לוד",
    "מודיעין-מכבים-רעות", "נהריה", "נוף הגליל", "כסייפה", "נצרת", "נשר", "נתיבות", "נתניה",
    "סח'נין", "עכו", "עפולה", "ערערה", "פרדס חנה-כרכור", "פתח תקווה", "צפת", "קריית אונו",
    "קריית אתא", "קריית ביאליק", "קריית גת", "קריית ים", "קריית מוצקין", "קריית מלאכי", "קריית שמונה",
    "ראש העין", "ראשון לציון", "רהט", "רחובות", "רכסים", "רמלה", "רמת גן", "רמת השרון", "רעננה",
    "שגב-שלום", "שדרות", "שפרעם", "תל אביב-יפו",
]

# Sizing: (base_sqm, inverter_kw, base_installed_kw)
SIZING_TIERS = [
    (80, 15, 18), (110, 20, 24), (150, 30, 36), (190, 40, 48), (240, 45, 54), (295, 50, 60),
]

HEBREW_MONTHS = [
    "ינואר", "פברואר", "מרץ", "אפריל", "מאי", "יוני",
    "יולי", "אוגוסט", "ספטמבר", "אוקטובר", "נובמבר", "דצמבר",
]

DEFAULT_CONFIG = {
    "representative": "רון הלל",
    "city": "",
    "roof_sqm": 100,
    "price_per_kw": 3100,
    "quote_number": 150,
}


# ---------------------------------------------------------------------------
# Config helpers
# ---------------------------------------------------------------------------

def load_config():
    try:
        with open(CONFIG_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
        cfg = DEFAULT_CONFIG.copy()
        for k in cfg:
            if k in data:
                v = data[k]
                if k == "roof_sqm":
                    cfg[k] = int(v) if v is not None else DEFAULT_CONFIG[k]
                elif k == "price_per_kw":
                    cfg[k] = float(v) if v is not None else DEFAULT_CONFIG[k]
                elif k == "quote_number":
                    cfg[k] = int(v) if v is not None else DEFAULT_CONFIG[k]
                else:
                    cfg[k] = v
        return cfg
    except (FileNotFoundError, json.JSONDecodeError, KeyError, ValueError, TypeError):
        return DEFAULT_CONFIG.copy()


def save_config(cfg):
    try:
        with open(CONFIG_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        data = {}
    data.update({k: v for k, v in cfg.items() if k in DEFAULT_CONFIG})
    with open(CONFIG_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


# ---------------------------------------------------------------------------
# Business logic helpers
# ---------------------------------------------------------------------------

def sanitize_folder_name(name):
    """Strip characters that are invalid in folder names."""
    name = name.strip()
    name = re.sub(r'[/\\:*?"<>|]', "_", name)
    return name or "לקוח"


def get_sizing(roof_sqm):
    if roof_sqm < 80:
        raise ValueError("גודל הגג חייב להיות לפחות 80 מ״ר")
    tier = None
    for base_sqm, inv_kw, base_inst in SIZING_TIERS:
        if roof_sqm >= base_sqm:
            tier = (base_sqm, inv_kw, base_inst)
    if not tier:
        raise ValueError("גודל הגג חייב להיות לפחות 80 מ״ר")
    base_sqm, inverter_kw, base_installed = tier
    extra_blocks = (roof_sqm - base_sqm) // 10
    installed_kw = base_installed + extra_blocks * 2
    max_installed = inverter_kw * 1.4
    installed_kw = min(installed_kw, max_installed)
    return inverter_kw, installed_kw


def cost_logic(installed_kw, price_per_kw):
    total_before_vat = installed_kw * price_per_kw
    vat = math.floor(total_before_vat * 0.18)
    total_after_vat = total_before_vat + vat
    return total_before_vat, vat, total_after_vat


def yearly_revenue(inverter_kw, urban_premium):
    rate_low = 0.54 if urban_premium else 0.48
    rate_high = 0.47 if urban_premium else 0.41
    if inverter_kw <= 15:
        rev = inverter_kw * 1650 * rate_low
    else:
        rev = 15 * 1650 * rate_low + (inverter_kw - 15) * 1650 * rate_high
    return rev


def blended_avg_per_kw(yearly_rev, inverter_kw):
    if inverter_kw <= 0:
        return "0.000"
    total_kwh = inverter_kw * 1650
    return "{:.3f}".format(yearly_rev / total_kwh)


def hebrew_date():
    now = datetime.now()
    return "{} ב{} {}".format(now.day, HEBREW_MONTHS[now.month - 1], now.year)


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _apply_document_font(doc, font_name="Times New Roman", size_pt=14):
    from docx.shared import Pt
    size = Pt(size_pt)
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = font_name
            run.font.size = size
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = font_name
                        run.font.size = size
    for section in doc.sections:
        for part in (
            section.header,
            section.footer,
            getattr(section, "first_page_header", None),
            getattr(section, "first_page_footer", None),
        ):
            if part is None:
                continue
            if hasattr(part, "paragraphs"):
                for p in part.paragraphs:
                    for run in p.runs:
                        run.font.name = font_name
                        run.font.size = size
            if hasattr(part, "tables"):
                for table in part.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.font.name = font_name
                                    run.font.size = size


def _cell_text(cell):
    parts = []
    for p in cell.paragraphs:
        for run in p.runs:
            if run.text:
                parts.append(run.text)
    return "".join(parts).strip()


def _set_cell_value(cell, value):
    for p in cell.paragraphs:
        p.clear()
    if not cell.paragraphs:
        cell.add_paragraph(str(value))
    else:
        cell.paragraphs[0].add_run(str(value))


# ---------------------------------------------------------------------------
# Letter template
# ---------------------------------------------------------------------------

def process_letter_template(output_folder_path, client_name, rep_full_name, rep_first_name):
    """
    Fill "לכבוד דייר.docx":
      - __________________ → client_name  (the "לכבוד" line)
      - rep placeholders   → rep_full / rep_first
    Save as "מכתב_פתיחה_[client_name].docx" in output_folder_path.
    """
    from docx import Document
    from docx.shared import Inches

    template_path = os.path.join(SOLRAY_ROOT, TEMPLATE_LETTER)
    if not os.path.isfile(template_path):
        raise FileNotFoundError("קובץ המקור 'לכבוד דייר.docx' חסר בתקיית מערכת לסלולר")

    doc = Document(template_path)

    # Logo in header
    logo_path = None
    for name in LOGO_NAMES:
        p = os.path.join(SOLRAY_ROOT, name)
        if os.path.isfile(p):
            logo_path = p
            break
    if logo_path:
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.clear()
        header_para.alignment = 0
        run = header_para.add_run()
        run.add_picture(logo_path, width=Inches(1.2))

    # Full-name replacement first so "רון הלל"→rep_full wins before "רון"→rep_first
    replacements = [
        ("__________________", client_name),
        ("רון הלל", rep_full_name),
        ("רון", rep_first_name),
    ]

    def replace_in_paragraph(paragraph):
        full_text = "".join(r.text or "" for r in paragraph.runs)
        new_text = full_text
        for old, new in replacements:
            new_text = new_text.replace(old, new)
        if new_text != full_text and paragraph.runs:
            paragraph.runs[0].text = new_text
            for r in paragraph.runs[1:]:
                r.text = ""

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)

    out_name = "מכתב_פתיחה_{}.docx".format(client_name)
    out_path = os.path.join(output_folder_path, out_name)
    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
# Quote template
# ---------------------------------------------------------------------------

def process_quote_template(
    output_folder_path,
    client_name,
    address,
    quote_number,
    roof_sqm,
    inverter_kw,
    installed_kw,
    price_per_kw,
    total_before_vat,
    vat,
    total_after_vat,
    city,
    urban_premium,
    blended_avg,
    yearly_rev,
    extra_items=None,
):
    """
    Fill "פורמט הצעת מחיר - מעוצב.docx".
    Template tables:
      0 – title bar (no changes)
      1 – 2×2 client/date/address/quote# (each cell: bold label + value paragraph)
      2 – 1×3 with nested tables: col0 = property data, col2 = income calculation
      3 – pricing: header row + solar data row + extra item rows + total row
      4 – signature (no changes)
    """
    import copy
    from docx import Document
    from docx.shared import Inches
    from docx.oxml.ns import qn
    from lxml import etree as _et

    template_path = os.path.join(SOLRAY_ROOT, TEMPLATE_QUOTE)
    if not os.path.isfile(template_path):
        raise FileNotFoundError("קובץ המקור '{}' חסר בתקיית מערכת לסלולר".format(TEMPLATE_QUOTE))

    doc = Document(template_path)

    # Logo in header
    logo_path = None
    for name in LOGO_NAMES:
        p = os.path.join(SOLRAY_ROOT, name)
        if os.path.isfile(p):
            logo_path = p
            break
    if logo_path:
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.clear()
        header_para.alignment = 0
        run = header_para.add_run()
        run.add_picture(logo_path, width=Inches(1.2))

    def fmt_num(x):
        return "{:,.0f}".format(round(x))

    _W_TBL = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl"
    _W_TR  = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr"
    _W_TC  = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc"
    _W_T   = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
    _W_P   = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
    _W_R   = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"

    def _set_nested_cell(tc_elem, value):
        """Set text of a raw <w:tc> element (handles nested table cells)."""
        t_nodes = list(tc_elem.iter(_W_T))
        if t_nodes:
            t_nodes[0].text = str(value)
            for t in t_nodes[1:]:
                t.text = ""
        else:
            p = tc_elem.find(_W_P)
            if p is None:
                p = _et.SubElement(tc_elem, _W_P)
            r = _et.SubElement(p, _W_R)
            t = _et.SubElement(r, _W_T)
            t.text = str(value)

    def _set_value_para(cell, value):
        """Set value in paragraph[1] of a cell (blank line after the bold label)."""
        para = cell.paragraphs[1] if len(cell.paragraphs) >= 2 else cell.paragraphs[0]
        para.clear()
        para.add_run(str(value))

    # --- Table 1: proposal details ---
    t1 = doc.tables[1]
    _set_value_para(t1.rows[0].cells[0], client_name)        # שם הלקוח
    _set_value_para(t1.rows[0].cells[1], hebrew_date())      # תאריך
    _set_value_para(t1.rows[1].cells[0], address)            # כתובת הנכס
    _set_value_para(t1.rows[1].cells[1], str(quote_number))  # מספר הצעה

    # --- Table 2: nested property + income tables ---
    t2 = doc.tables[2]
    t2_tcs = t2.rows[0]._tr.findall(_W_TC)

    prop_tbl = t2_tcs[0].find(_W_TBL)
    if prop_tbl is not None:
        prop_rows = prop_tbl.findall(_W_TR)
        # row0 = header "נתוני הנכס", rows 1-4 = data
        for i, val in enumerate([address, str(roof_sqm), str(inverter_kw), str(installed_kw)]):
            ri = i + 1
            if ri < len(prop_rows):
                tc_list = prop_rows[ri].findall(_W_TC)
                if len(tc_list) >= 2:
                    _set_nested_cell(tc_list[1], val)

    inc_tbl = t2_tcs[2].find(_W_TBL)
    if inc_tbl is not None:
        inc_rows = inc_tbl.findall(_W_TR)
        # row0 = header, row1=עיר, row2=פרמיה, row3=תמחור ממוצע, row4=שעות (skip), row5=הכנסה
        for ri, val in {
            1: city,
            2: "כן" if urban_premium else "לא",
            3: blended_avg,
            5: "₪ " + fmt_num(yearly_rev),
        }.items():
            if ri < len(inc_rows):
                tc_list = inc_rows[ri].findall(_W_TC)
                if len(tc_list) >= 2:
                    _set_nested_cell(tc_list[1], val)

    # --- Table 3: pricing ---
    t3 = doc.tables[3]
    solar_total = installed_kw * price_per_kw

    # Row 1: solar system (col0 = description label already in template)
    solar_cells = t3.rows[1].cells
    _set_cell_value(solar_cells[1], "₪ " + fmt_num(price_per_kw))
    _set_cell_value(solar_cells[2], str(installed_kw))
    _set_cell_value(solar_cells[3], "₪ " + fmt_num(solar_total))

    # Extra item rows inserted after solar row, before total row
    for i, item in enumerate(extra_items or []):
        src_row = t3.rows[1 + i]
        new_tr = copy.deepcopy(src_row._tr)
        for t_elem in new_tr.iter(qn("w:t")):
            t_elem.text = ""
        src_row._tr.addnext(new_tr)
        new_cells = t3.rows[2 + i].cells
        _set_cell_value(new_cells[0], item["desc"])
        _set_cell_value(new_cells[1], "₪ " + fmt_num(item["price"]))
        qty_val = item["qty"]
        _set_cell_value(new_cells[2], str(int(qty_val)) if qty_val == int(qty_val) else str(qty_val))
        _set_cell_value(new_cells[3], "₪ " + fmt_num(item["total"]))

    # Total row (last row, orange): col3 = total before VAT
    _set_cell_value(t3.rows[-1].cells[3], "₪ " + fmt_num(total_before_vat))

    out_name = "הצעת מחיר {}.docx".format(client_name)
    out_path = os.path.join(output_folder_path, out_name)
    doc.save(out_path)
    return out_path


def merge_letter_and_quote(letter_path, quote_path, output_path):
    """
    Single document: first page(s) = letter, then page break, then quote.
    Uses docxcompose so images and relationships stay valid.
    """
    from docx import Document
    from docxcompose.composer import Composer

    master = Document(letter_path)
    master.add_page_break()
    composer = Composer(master)
    composer.append(Document(quote_path))
    composer.save(output_path)


# ---------------------------------------------------------------------------
# Flask app
# ---------------------------------------------------------------------------

def create_app():
    from flask import Flask, render_template_string, request, redirect, url_for, send_file

    app = Flask(__name__)
    app.secret_key = os.urandom(16).hex()
    app.config["JSON_AS_ASCII"] = False

    HTML_TEMPLATE = r"""<!DOCTYPE html>
<html dir="rtl" lang="he">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>סול-ריי – הצעת מחיר</title>
  <style>
    *, *::before, *::after { box-sizing: border-box; }
    body { font-family: system-ui, 'Segoe UI', sans-serif; max-width: 660px; margin: 2rem auto; padding: 0 1rem; background: #f5f7fa; color: #1a1a2e; }
    h1 { font-size: 1.5rem; margin-bottom: 1.2rem; }
    h2 { font-size: 1rem; margin: 0 0 0.8rem; color: #e85d04; border-bottom: 2px solid #e85d04; padding-bottom: 0.3rem; }
    label { display: block; margin-top: 0.75rem; font-weight: 600; font-size: 0.92rem; }
    select, input[type="text"], input[type="number"] { width: 100%; padding: 0.45rem 0.6rem; border: 1px solid #ccc; border-radius: 6px; margin-top: 0.2rem; font-size: 0.93rem; background: #fff; }
    select:focus, input:focus { outline: none; border-color: #e85d04; box-shadow: 0 0 0 2px rgba(232,93,4,0.15); }
    .field { margin-bottom: 0.9rem; }
    .hint { font-size: 0.8rem; color: #777; margin-top: 0.25rem; }
    .suggest-link { font-size: 0.82rem; color: #e85d04; cursor: pointer; text-decoration: underline; margin-top: 0.3rem; display: inline-block; }
    .suggest-link:hover { color: #c84d00; }
    .error { color: #b00; background: #fee8e8; padding: 0.65rem 1rem; margin: 1rem 0; border-radius: 6px; border-right: 3px solid #b00; }
    .success-box { background: #fff; border: 1.5px solid #166534; border-radius: 10px; padding: 1rem 1.3rem; margin: 1rem 0; }
    .success-box p { color: #166534; margin: 0 0 0.8rem; font-weight: 600; }
    .dl-btn { display: inline-block; padding: 0.55rem 1.2rem; border-radius: 7px; font-size: 0.95rem; font-weight: 700; text-decoration: none; margin-left: 0.5rem; margin-bottom: 0.4rem; }
    .dl-letter { background: #e8f5e9; color: #166534; border: 1.5px solid #4caf50; }
    .dl-quote  { background: #fff3e8; color: #c84d00; border: 1.5px solid #e85d04; }
    .dl-btn:hover { filter: brightness(0.92); }
    .form-card { background: #fff; border-radius: 10px; padding: 1.1rem 1.4rem 1.3rem; margin-bottom: 1rem; box-shadow: 0 2px 8px rgba(0,0,0,0.07); }
    .preview-card { background: #fff; border: 1.5px solid #e85d04; border-radius: 10px; padding: 1rem 1.3rem 1.1rem; margin: 1rem 0; box-shadow: 0 2px 10px rgba(232,93,4,0.08); }
    .preview-card h3 { margin: 0 0 0.7rem; font-size: 0.95rem; color: #e85d04; }
    .preview-table { width: 100%; border-collapse: collapse; font-size: 0.9rem; }
    .preview-table td { padding: 0.28rem 0.15rem; }
    .preview-table td:last-child { text-align: left; font-variant-numeric: tabular-nums; white-space: nowrap; }
    .preview-table .sys-row td { color: #444; }
    .preview-table tr.sep td { border-top: 1px solid #e0e0e0; padding-top: 0.45rem; }
    .preview-table tr.total-row td { font-weight: 700; font-size: 1rem; border-top: 2px solid #e85d04; padding-top: 0.5rem; color: #1a1a2e; }
    .extra-header { display: flex; gap: 0.4rem; margin-bottom: 0.3rem; font-size: 0.78rem; color: #999; font-weight: 600; padding: 0 0.2rem; }
    .extra-header .eh-desc { flex: 2.5; }
    .extra-header .eh-qty, .extra-header .eh-price { flex: 1.2; }
    .extra-header .eh-del { width: 28px; }
    .extra-row { display: flex; gap: 0.4rem; margin-bottom: 0.45rem; align-items: center; }
    .extra-row input[type="text"] { flex: 2.5; min-width: 0; }
    .extra-row input[type="number"] { flex: 1.2; min-width: 0; }
    .del-btn { background: #fee; border: 1px solid #fbb; color: #c00; border-radius: 5px; padding: 0.28rem 0.55rem; cursor: pointer; font-size: 0.85rem; flex-shrink: 0; line-height: 1; }
    .del-btn:hover { background: #fdd; }
    .add-btn { background: #fff3e8; border: 1px solid #e85d04; color: #e85d04; border-radius: 6px; padding: 0.38rem 1rem; cursor: pointer; font-size: 0.88rem; margin-top: 0.3rem; }
    .add-btn:hover { background: #ffe0c8; }
    .submit-btn { display: block; width: 100%; margin-top: 1.4rem; padding: 0.78rem; font-size: 1.05rem; font-weight: 700; background: #e85d04; color: #fff; border: none; border-radius: 8px; cursor: pointer; letter-spacing: 0.02em; }
    .submit-btn:hover { background: #c84d00; }
  </style>
</head>
<body>
  <h1>סול-ריי – הצעת מחיר</h1>
  {% if error %}<div class="error">{{ error }}</div>{% endif %}
  {% if document_url %}
  <div class="success-box">
    <p>המסמך נוצר בהצלחה! (הצעה מס׳ {{ quote_number }}) — דף ראשון: מכתב פתיחה, דף שני: הצעת מחיר</p>
    <a class="dl-btn dl-letter" href="{{ document_url }}" download>הורד מסמך Word</a>
  </div>
  {% endif %}
  <form method="post" action="{{ url_for('generate') }}">

    <div class="form-card">
      <h2>פרטי הנכס</h2>
      <div class="field">
        <label for="representative">נציג</label>
        <select id="representative" name="representative" required>
          {% for rep in representatives %}
          <option value="{{ rep.full }}" {{ 'selected' if rep.full == config.representative else '' }}>{{ rep.full }}</option>
          {% endfor %}
        </select>
      </div>
      <div class="field">
        <label for="client_name">שם לקוח</label>
        <input type="text" id="client_name" name="client_name" value="{{ prefill_client }}" placeholder="ישראל ישראלי" required>
        <p class="hint">יופיע אחרי ״לכבוד״ במכתב ובכותרת ההצעה</p>
      </div>
      <div class="field">
        <label for="address">כתובת</label>
        <input type="text" id="address" name="address" value="{{ prefill_address }}" placeholder="הרצל 10, תל אביב" required>
        <p class="hint">תכנס לטבלת נתוני הנכס בהצעת המחיר</p>
      </div>
      <div class="field">
        <label for="city">עיר</label>
        <input type="text" id="city" name="city" list="cities_list" value="{{ config.city or '' }}" placeholder="הקלד או בחר עיר" autocomplete="off">
        <datalist id="cities_list">
          {% for c in cities %}<option value="{{ c }}">{% endfor %}
        </datalist>
      </div>
    </div>

    <div class="form-card">
      <h2>נתוני מערכת</h2>
      <div class="field">
        <label for="roof_sqm">גודל הגג (מ״ר)</label>
        <input type="number" id="roof_sqm" name="roof_sqm" min="80" step="1" value="{{ config.roof_sqm }}" required oninput="updatePreview()">
        <p class="hint">מינימום 80 מ״ר</p>
      </div>
      <div class="field">
        <label for="price_per_kw">מחיר לקילוואט (₪)</label>
        <input type="number" id="price_per_kw" name="price_per_kw" min="0" step="1" value="{{ config.price_per_kw }}" required oninput="updatePreview()">
        <span class="suggest-link" id="price-suggest" style="display:none" onclick="applySuggestedPrice()"></span>
      </div>
    </div>

    <div id="preview"></div>

    <div class="form-card">
      <h2>פריטים נוספים לתמחור</h2>
      <div class="extra-header">
        <span class="eh-desc">פירוט</span>
        <span class="eh-qty">כמות</span>
        <span class="eh-price">מחיר ליחידה (₪)</span>
        <span class="eh-del"></span>
      </div>
      <div id="extra-items-container"></div>
      <button type="button" class="add-btn" onclick="addExtraItem()">+ הוסף פריט</button>
    </div>

    <button type="submit" class="submit-btn">הפק מסמכים</button>
  </form>

  <script>
    const SIZING_TIERS = [[80,15,18],[110,20,24],[150,30,36],[190,40,48],[240,45,54],[295,50,60]];
    let _suggestedPrice = 0;

    function getSizing(sqm) {
      if (sqm < 80) return null;
      let tier = null;
      for (const t of SIZING_TIERS) { if (sqm >= t[0]) tier = t; }
      if (!tier) return null;
      const [baseSqm, invKw, baseInst] = tier;
      const extra = Math.floor((sqm - baseSqm) / 10);
      const inst = Math.min(baseInst + extra * 2, invKw * 1.4);
      return { inverter: invKw, installed: inst };
    }

    function suggestedPriceFor(invKw) {
      if (invKw <= 15) return 3100;
      if (invKw <= 50) return 2600;
      return 2400;
    }

    function fmtILS(n) {
      return '\u20aa\u202f' + Math.round(n).toLocaleString('he-IL');
    }

    function applySuggestedPrice() {
      if (_suggestedPrice > 0) {
        document.getElementById('price_per_kw').value = _suggestedPrice;
        updatePreview();
      }
    }

    function updatePreview() {
      const sqm = parseInt(document.getElementById('roof_sqm').value) || 0;
      const ppkw = parseFloat(document.getElementById('price_per_kw').value) || 0;
      const previewDiv = document.getElementById('preview');
      const suggestSpan = document.getElementById('price-suggest');

      if (sqm < 80) { previewDiv.innerHTML = ''; suggestSpan.style.display = 'none'; return; }
      const sizing = getSizing(sqm);
      if (!sizing) { previewDiv.innerHTML = ''; suggestSpan.style.display = 'none'; return; }

      _suggestedPrice = suggestedPriceFor(sizing.inverter);
      suggestSpan.textContent = 'מחיר מוצע לפי טבלה: ' + fmtILS(_suggestedPrice) + ' לקו״א (לחץ להחלה)';
      suggestSpan.style.display = 'inline-block';

      let extraTotal = 0;
      document.querySelectorAll('.extra-row').forEach(row => {
        const qty = parseFloat(row.querySelector('.extra-qty').value) || 0;
        const price = parseFloat(row.querySelector('.extra-price').value) || 0;
        extraTotal += qty * price;
      });

      const solarTotal = sizing.installed * ppkw;
      const combinedBefore = solarTotal + extraTotal;
      const vatAmt = Math.floor(combinedBefore * 0.18);
      const grandTotal = combinedBefore + vatAmt;

      let rows = `
        <tr class="sys-row"><td>גודל ממיר</td><td><strong>${sizing.inverter} קוו״א</strong></td></tr>
        <tr class="sys-row"><td>כמות קילוואט להתקנה</td><td><strong>${sizing.installed} קוו״א</strong></td></tr>`;
      if (ppkw > 0) {
        rows += `<tr><td>עלות מערכת סולארית</td><td>${fmtILS(solarTotal)}</td></tr>`;
        if (extraTotal > 0) rows += `<tr><td>פריטים נוספים</td><td>${fmtILS(extraTotal)}</td></tr>`;
        rows += `
        <tr class="sep"><td>סה״כ לפני מע״מ</td><td>${fmtILS(combinedBefore)}</td></tr>
        <tr><td>מע״מ 18%</td><td>${fmtILS(vatAmt)}</td></tr>
        <tr class="total-row"><td>סה״כ לתשלום</td><td>${fmtILS(grandTotal)}</td></tr>`;
      }
      previewDiv.innerHTML = `<div class="preview-card"><h3>תצוגה מקדימה</h3><table class="preview-table">${rows}</table></div>`;
    }

    function addExtraItem() {
      const container = document.getElementById('extra-items-container');
      const row = document.createElement('div');
      row.className = 'extra-row';
      row.innerHTML =
        '<input type="text" name="extra_desc" placeholder="פירוט" oninput="updatePreview()">' +
        '<input type="number" name="extra_qty" class="extra-qty" min="0" step="0.01" placeholder="כמות" oninput="updatePreview()">' +
        '<input type="number" name="extra_price" class="extra-price" min="0" step="0.01" placeholder="מחיר" oninput="updatePreview()">' +
        '<button type="button" class="del-btn" onclick="this.parentElement.remove();updatePreview()">✕</button>';
      container.appendChild(row);
    }

    window.addEventListener('DOMContentLoaded', updatePreview);
  </script>
</body>
</html>"""

    @app.route("/")
    def index():
        config = load_config()
        return render_template_string(
            HTML_TEMPLATE,
            config=config,
            representatives=REPRESENTATIVES,
            cities=CITIES,
            prefill_client=request.args.get("client", ""),
            prefill_address=request.args.get("address", ""),
            error=request.args.get("error"),
            document_url=request.args.get("document_url"),
            quote_number=request.args.get("quote_number", ""),
        )

    @app.route("/generate", methods=["POST"])
    def generate():
        client_name = (request.form.get("client_name") or "").strip()
        address = (request.form.get("address") or "").strip()

        if not client_name:
            return redirect(url_for("index", error="נא להזין שם לקוח."))
        if not address:
            return redirect(url_for("index", error="נא להזין כתובת."))

        if not os.path.isfile(os.path.join(SOLRAY_ROOT, TEMPLATE_LETTER)):
            return redirect(url_for("index", error="קובץ 'לכבוד דייר.docx' חסר בתקיית מערכת לסלולר."))
        if not os.path.isfile(os.path.join(SOLRAY_ROOT, TEMPLATE_QUOTE)):
            return redirect(url_for("index", error="קובץ '{}' חסר בתקיית מערכת לסלולר.".format(TEMPLATE_QUOTE)))

        try:
            roof_sqm = int(request.form.get("roof_sqm") or "0")
        except (TypeError, ValueError):
            return redirect(url_for("index", error="גודל הגג חייב להיות מספר שלם."))
        if roof_sqm < 80:
            return redirect(url_for("index", error="גודל הגג חייב להיות לפחות 80 מ״ר."))

        try:
            price_per_kw = float(request.form.get("price_per_kw") or "0")
        except (TypeError, ValueError):
            return redirect(url_for("index", error="מחיר לקילוואט חייב להיות מספר."))
        if price_per_kw <= 0:
            return redirect(url_for("index", error="מחיר לקילוואט חייב להיות חיובי."))

        city = (request.form.get("city") or "").strip()
        representative = request.form.get("representative") or "רון הלל"
        rep_first = next(
            (r["first"] for r in REPRESENTATIVES if r["full"] == representative),
            representative.split()[0] if representative else "רון",
        )

        config = load_config()
        quote_number = int(config.get("quote_number", 1))
        urban_premium = city in CITIES

        try:
            inverter_kw, installed_kw = get_sizing(roof_sqm)
        except ValueError as e:
            return redirect(url_for("index", error=str(e)))

        extra_descs = request.form.getlist("extra_desc")
        extra_qtys = request.form.getlist("extra_qty")
        extra_prices = request.form.getlist("extra_price")
        extra_items = []
        for i in range(len(extra_descs)):
            desc = extra_descs[i].strip() if i < len(extra_descs) else ""
            try:
                qty = float(extra_qtys[i]) if i < len(extra_qtys) and extra_qtys[i].strip() else 0.0
            except (ValueError, IndexError):
                qty = 0.0
            try:
                eprice = float(extra_prices[i]) if i < len(extra_prices) and extra_prices[i].strip() else 0.0
            except (ValueError, IndexError):
                eprice = 0.0
            if desc or qty or eprice:
                extra_items.append({
                    "desc": desc or "פריט נוסף",
                    "qty": qty,
                    "price": eprice,
                    "total": qty * eprice,
                })

        solar_before_vat, _sv, _sa = cost_logic(installed_kw, price_per_kw)
        extra_total = sum(item["total"] for item in extra_items)
        combined_before_vat = solar_before_vat + extra_total
        vat = math.floor(combined_before_vat * 0.18)
        total_after_vat = combined_before_vat + vat
        total_before_vat = combined_before_vat

        yearly_rev = yearly_revenue(inverter_kw, urban_premium)
        blended_avg = blended_avg_per_kw(yearly_rev, inverter_kw)

        # Create output folder: לקוחות/<sanitized_client_name>/
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        client_folder = os.path.join(OUTPUT_DIR, sanitize_folder_name(client_name))
        os.makedirs(client_folder, exist_ok=True)

        try:
            letter_path = process_letter_template(client_folder, client_name, representative, rep_first)
            quote_path = process_quote_template(
                client_folder,
                client_name,
                address,
                quote_number,
                roof_sqm,
                inverter_kw,
                installed_kw,
                price_per_kw,
                total_before_vat,
                vat,
                total_after_vat,
                city,
                urban_premium,
                blended_avg,
                yearly_rev,
                extra_items=extra_items,
            )
        except FileNotFoundError as e:
            return redirect(url_for("index", error=str(e)))
        except Exception as e:
            return redirect(url_for("index", error="שגיאה בעיבוד: {}".format(e)))

        combined_name = "מכתב_פתיחה_והצעת_מחיר_{}.docx".format(client_name)
        combined_path = os.path.join(client_folder, combined_name)
        try:
            merge_letter_and_quote(letter_path, quote_path, combined_path)
        except Exception as e:
            return redirect(url_for("index", error="שגיאה באיחוד המסמכים: {}".format(e)))

        try:
            os.remove(letter_path)
            os.remove(quote_path)
        except OSError:
            pass

        quote_number += 1
        save_config({
            "representative": representative,
            "city": city,
            "roof_sqm": roof_sqm,
            "price_per_kw": price_per_kw,
            "quote_number": quote_number,
        })

        combined_rel = os.path.relpath(combined_path, OUTPUT_DIR)

        return redirect(url_for(
            "index",
            client=client_name,
            address=address,
            document_url=url_for("download", filename=combined_rel),
            quote_number=quote_number - 1,
        ))

    @app.route("/download/<path:filename>")
    def download(filename):
        output_abs = os.path.abspath(OUTPUT_DIR)
        file_path = os.path.abspath(os.path.join(OUTPUT_DIR, filename))
        if not file_path.startswith(output_abs + os.sep):
            return "Forbidden", 403
        if not os.path.isfile(file_path):
            return "File not found", 404
        return send_file(file_path, as_attachment=True)

    return app


if __name__ == "__main__":
    import threading
    import webbrowser

    app = create_app()
    port = 5051
    url = "http://127.0.0.1:{}/".format(port)

    def open_browser():
        webbrowser.open(url)

    threading.Timer(1.0, open_browser).start()
    app.run(host="0.0.0.0", port=port, debug=False, use_reloader=False)
